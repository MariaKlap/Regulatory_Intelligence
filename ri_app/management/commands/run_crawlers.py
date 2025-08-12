import os
import glob
import time
import logging
import tempfile
import requests
import subprocess
import pandas as pd
import sqlite3
from django.core.management.base import BaseCommand
from django.conf import settings
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta
from ...models import RegulatoryData

logger = logging.getLogger(__name__)

class Command(BaseCommand):
    help = 'Run all regulatory intelligence crawlers and update database'
    
    # List of raw GitHub URLs to your scripts
    GITHUB_SCRIPTS = [
        "https://raw.githubusercontent.com/MariaKlap/RI/main/EMAnews2.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/ECnews11.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/ICR.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/ICHnews.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/IS1.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/SWISS5.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/AT.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/GMP.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/EC-Updates.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/EC-Medical.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/FDAnews.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/RQAnews4.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/Topra.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/raps-2.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/WHOnews.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/CBGnewsfinal5win.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/main/HMA6news.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/BEnews1.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/CY.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/DE.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/DK3newswin.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/FInew.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/IE.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/Infarmed6news.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/Luxnews.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/MHRA.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/MHRANews.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/MHRAPolicy.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/Maltanews.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/Norwnews%20(2).py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/SEn.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/SEns.py",
        "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/SEnsa.py",
    ]

    def add_arguments(self, parser):
        parser.add_argument(
            '--force',
            action='store_true',
            help='Force run all crawlers even if recent data exists'
        )
        parser.add_argument(
            '--skip-docx',
            action='store_true',
            help='Skip generating the DOCX report'
        )
        parser.add_argument(
            '--cleanup',
            action='store_true',
            help='Clean up temporary files after run'
        )
        parser.add_argument(
            '--keep-old-data',
            action='store_true',
            help='Keep existing data in database'
        )

    def handle(self, *args, **options):
        # Set up logging
        log_file = os.path.join(settings.BASE_DIR, "batch_run_log.txt")
        logging.basicConfig(
            filename=log_file,
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        
        self.stdout.write(self.style.SUCCESS("=== Starting Regulatory Intelligence Data Collection ==="))
        logging.info("=== Batch GitHub Execution Started ===")
        
        try:
            # Run all crawlers
            for url in self.GITHUB_SCRIPTS:
                self.download_and_run_script(url)
            
            # Combine and process data
            if self.combine_excel_files():
                if self.convert_excel_to_db():
                    if self.compare_with_github_xlsx():
                        if not options['skip_docx']:
                            self.export_news_to_docx()
            
            # Import data to Django models
            self.import_to_django(options['keep_old_data'])
            
            # Cleanup if requested
            if options['cleanup']:
                self.cleanup_temp_files()
            
            self.stdout.write(self.style.SUCCESS("=== Successfully Completed Data Collection ==="))
            logging.info("=== Batch GitHub Execution Completed ===")
        
        except Exception as e:
            logger.error(f"Critical error in run_crawlers: {str(e)}")
            self.stdout.write(self.style.ERROR(f"Error: {str(e)}"))

    def download_and_run_script(self, url):
        try:
            logging.info(f"📥 Downloading script: {url}")
            self.stdout.write(f"Downloading {url.split('/')[-1]}...")
            
            response = requests.get(url, timeout=30)
            response.raise_for_status()

            with tempfile.NamedTemporaryFile(delete=False, suffix=".py") as temp_file:
                temp_file.write(response.content)
                temp_file_path = temp_file.name

            logging.info(f"🚀 Running script: {url}")
            subprocess.run(["python", temp_file_path], check=True)
            logging.info(f"✅ Completed: {url}")
            
            os.remove(temp_file_path)
            return True
            
        except requests.RequestException as e:
            logging.error(f"❌ Download failed for {url}: {e}")
            self.stdout.write(self.style.WARNING(f"Download failed for {url}"))
            return False
        except subprocess.CalledProcessError as e:
            logging.error(f"❌ Execution failed for {url}: {e}")
            self.stdout.write(self.style.WARNING(f"Execution failed for {url}"))
            return False
        except Exception as e:
            logging.error(f"❌ Unexpected error with {url}: {e}")
            return False

    def combine_excel_files(self):
        """Combine all Excel files in the directory into one RI.xlsx"""
        try:
            logging.info("🔍 Searching for Excel files to combine...")
            self.stdout.write("Combining Excel files...")
            
            excel_files = glob.glob(os.path.join(settings.BASE_DIR, '*.xlsx'))
            excel_files = [f for f in excel_files if not f.endswith('RI.xlsx')]

            if not excel_files:
                logging.warning("⚠️ No Excel files found to combine")
                self.stdout.write(self.style.WARNING("No Excel files found to combine"))
                return False

            combined_df = pd.DataFrame()
            
            for file in excel_files:
                try:
                    df = pd.read_excel(file, keep_default_na=True)
                    df['Source_File'] = os.path.basename(file)
                    combined_df = pd.concat([combined_df, df], ignore_index=True)
                    logging.info(f"➕ Added {file} to combined dataframe")
                except Exception as e:
                    logging.error(f"❌ Error reading {file}: {e}")
                    continue

            # Enhanced Date Processing
            if 'Date' in combined_df.columns:
                try:
                    # Step 1: Clean date strings
                    combined_df['Date'] = combined_df['Date'].replace(
                        ['None', 'N/A', 'NA', '', 'NaN', 'nan'], pd.NA
                    )
                    
                    # Step 2: Standardize separators
                    combined_df['Date'] = combined_df['Date'].astype(str).str.replace(
                        r'[.-]', '/', regex=True
                    )
                    
                    # Step 3: Parse dates with multiple formats
                    parsed_dates = pd.to_datetime(
                        combined_df['Date'],
                        errors='coerce',
                        dayfirst=True,
                        format='mixed'
                    )
                    
                    combined_df['Date'] = parsed_dates
                    
                    # Step 4: Filter by date (keep last 12 months or None)
                    one_year_ago = pd.Timestamp.now() - pd.DateOffset(months=12)
                    initial_count = len(combined_df)
                    
                    combined_df = combined_df[
                        (combined_df['Date'].isna()) | 
                        (combined_df['Date'] >= one_year_ago)
                    ]
                    
                    filtered_count = initial_count - len(combined_df)
                    logging.info(f"🧹 Filtered out {filtered_count} records older than {one_year_ago.date()}")
                    
                    # Step 5: Format for output (preserve NaT as None)
                    combined_df['Date'] = combined_df['Date'].dt.strftime('%Y-%m-%d').where(
                        combined_df['Date'].notna(), None
                    )
                except Exception as e:
                    logging.error(f"❌ Failed to process 'Date' column: {e}")
                    combined_df['Date'] = combined_df['Date'].astype(str)

            if not combined_df.empty:
                output_path = os.path.join(settings.BASE_DIR, 'RI.xlsx')
                combined_df.to_excel(output_path, index=False, na_rep='None')
                logging.info(f"💾 Saved combined Excel to {output_path}")
                self.stdout.write(self.style.SUCCESS(f"Combined data saved to RI.xlsx"))
                return True
            else:
                logging.warning("⚠️ No data to save - combined dataframe is empty")
                self.stdout.write(self.style.WARNING("No data found in Excel files"))
                return False

        except Exception as e:
            logging.error(f"❌ Error combining Excel files: {e}")
            self.stdout.write(self.style.ERROR(f"Error combining files: {str(e)}"))
            return False

    def convert_excel_to_db(self):
        """Convert RI.xlsx to RI.db SQLite database and RI.csv file"""
        try:
            excel_path = os.path.join(settings.BASE_DIR, 'RI.xlsx')
            csv_path = os.path.join(settings.BASE_DIR, 'RI.csv')
            
            if not os.path.exists(excel_path):
                logging.warning("⚠️ RI.xlsx not found")
                self.stdout.write(self.style.WARNING("RI.xlsx not found"))
                return False
            
            logging.info("📊 Reading combined Excel file...")
            df = pd.read_excel(excel_path, keep_default_na=True)
            
            if df.empty:
                logging.warning("⚠️ Excel file is empty")
                self.stdout.write(self.style.WARNING("Excel file is empty"))
                return False
            
            # Save as CSV
            logging.info("💾 Creating CSV file...")
            df.to_csv(csv_path, index=False, na_rep='None')
            logging.info(f"✅ Successfully created CSV file at {csv_path}")
            self.stdout.write(self.style.SUCCESS("Created RI.csv"))
            
            return True

        except Exception as e:
            logging.error(f"❌ Error creating CSV: {e}")
            self.stdout.write(self.style.ERROR(f"Error creating CSV: {str(e)}"))
            return False
        
    def compare_with_github_csv(self):
        """
        Compare local RI.csv with GitHub RI.csv based on 'Article URL'.
        Write unmatched articles to News.xlsx.
        """
        try:
            local_csv_path = os.path.join(os.getcwd(), 'RI.csv')
            github_csv_url = "https://raw.githubusercontent.com/MariaKlap/Master-Script/refs/heads/main/RI.csv"
            output_excel_path = os.path.join(os.getcwd(), 'News.xlsx')
            
            if not os.path.exists(local_csv_path):
                logging.warning("⚠️ Local RI.csv not found. Skipping comparison.")
                return False

            # Load local CSV (assumed to be comma-separated)
            df_local = pd.read_csv(local_csv_path, keep_default_na=True)

            # Load GitHub CSV (semicolon-separated)
            df_github = pd.read_csv(
                github_csv_url,
                sep=';',
                keep_default_na=True,
                on_bad_lines='skip',  # if using Python < 3.10, replace with error_bad_lines=False
                quoting=1             # handle quoted fields properly
            )

            # Ensure required columns exist
            if 'Article URL' not in df_local.columns or 'Article URL' not in df_github.columns:
                logging.error("❌ 'Article URL' column missing")
                self.stdout.write(self.style.ERROR("Missing 'Article URL' column"))
                return False

            # Find new articles
            unmatched_df = df_local[~df_local['Article URL'].isin(df_github['Article URL'])]

            if unmatched_df.empty:
                logging.info("✅ No new articles found")
                self.stdout.write(self.style.SUCCESS("No new articles found"))
                return True

            # Prepare output
            required_columns = ['Title', 'Summary', 'Date', 'Article URL', 'Source_File']
            for col in required_columns:
                if col not in unmatched_df.columns:
                    unmatched_df[col] = None

            result_df = unmatched_df[required_columns]
            result_df.to_excel(output_excel_path, index=False)
            logging.info(f"📝 Saved new articles to {output_excel_path}")
            self.stdout.write(self.style.SUCCESS(f"Found {len(result_df)} new articles"))
            return True

        except Exception as e:
            logging.error(f"❌ Error during comparison: {e}")
            self.stdout.write(self.style.ERROR(f"Comparison error: {str(e)}"))
            return False

    def export_news_to_docx(self):
        """Convert News.xlsx into a formatted RI_News.docx"""
        try:
            excel_path = os.path.join(settings.BASE_DIR, 'News.xlsx')
            docx_path = os.path.join(settings.BASE_DIR, 'RI_News.docx')

            if not os.path.exists(excel_path):
                logging.warning("⚠️ News.xlsx not found")
                self.stdout.write(self.style.WARNING("News.xlsx not found"))
                return False

            df = pd.read_excel(excel_path)

            if df.empty:
                logging.info("⚠️ News.xlsx is empty")
                self.stdout.write(self.style.WARNING("No new articles to export"))
                return False

            doc = Document()
            doc.add_heading("Regulatory Intelligence News Report", level=0)
            
            # Add report metadata
            doc.add_paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"New articles found: {len(df)}")
            doc.add_paragraph()
            
            # Add table of contents
            doc.add_heading("Table of Contents", level=1)
            toc = doc.add_paragraph()
            run = toc.add_run()
            run.add_break()
            
            for idx, row in df.iterrows():
                title = str(row.get("Title", "")).strip()
                toc.add_run(f"{idx+1}. {title[:100]}...")
                toc.add_run().add_break()
            
            doc.add_page_break()
            
            # Add articles
            for idx, row in df.iterrows():
                title = str(row.get("Title", "")).strip()
                summary = str(row.get("Summary", "")).strip()
                date = str(row.get("Date", "")).strip()
                url = str(row.get("Article URL", "")).strip()
                source = str(row.get("Source_File", "")).strip()
                
                # Article heading
                doc.add_heading(f"Article {idx+1}: {title}", level=1)
                
                # Metadata paragraph
                meta = doc.add_paragraph()
                meta.add_run("Source: ").bold = True
                meta.add_run(f"{source}\t")
                
                meta.add_run("Date: ").bold = True
                meta.add_run(f"{date}\t")
                
                # Add hyperlink
                para = doc.add_paragraph()
                para.add_run("Link: ").bold = True
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), doc.part.relate_to(url, 
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", 
                    is_external=True))
                
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'Hyperlink')
                rPr.append(rStyle)
                new_run.append(rPr)
                
                text_elem = OxmlElement('w:t')
                text_elem.text = "View original article"
                new_run.append(text_elem)
                hyperlink.append(new_run)
                para._p.append(hyperlink)
                
                # Add summary
                doc.add_heading("Summary", level=2)
                doc.add_paragraph(summary)
                
                doc.add_page_break()
            
            doc.save(docx_path)
            logging.info(f"📝 Exported news to {docx_path}")
            self.stdout.write(self.style.SUCCESS(f"Report generated: RI_News.docx"))
            return True

        except Exception as e:
            logging.error(f"❌ Error exporting to DOCX: {e}")
            self.stdout.write(self.style.ERROR(f"DOCX export error: {str(e)}"))
            return False

    def import_to_django(self, keep_old_data=False):
        """Import data from RI.csv to Django models"""
        try:
            
            excel_path = os.path.join(settings.BASE_DIR, 'RI.xlsx')
            
            if not os.path.exists(excel_path):
                logging.warning("⚠️ RI.xlsx not found")
                self.stdout.write(self.style.WARNING("RI.csv not found - skipping import"))
                return False
            
            logging.info("Importing data to Django models...")
            self.stdout.write("Importing data to database...")
            
            df = pd.read_excel(excel_path, keep_default_na=True)

            if df.empty:
                logging.warning("⚠️ No data to import")
                self.stdout.write(self.style.WARNING("No data to import"))
                return False
            
            # Clear existing data unless keep_old_data is True
            if not keep_old_data:
                RegulatoryData.objects.all().delete()
                logging.info("Cleared existing data from database")
            
            # Prepare data for bulk create
            records = []
            existing_urls = set(RegulatoryData.objects.values_list('article_url', flat=True))
            
            for _, row in df.iterrows():
                url = str(row.get('Article URL', '')).strip()
                if not url or url in existing_urls:
                    continue
                
                try:
                    records.append(RegulatoryData(
                        title=str(row.get('Title', '')),
                        summary=str(row.get('Summary', '')),
                        date=self.parse_date(row.get('Date', '')),
                        article_url=url,
                        source_file=row.get('Source_File', ''),
                        agency=self.determine_agency(row.get('Source_File', '')),
                        category=self.determine_category(row.get('Source_File', ''))
                    ))
                except Exception as e:
                    logging.error(f"❌ Error processing row: {e}")
                    continue
            
            # Bulk create
            if records:
                RegulatoryData.objects.bulk_create(records)
                logging.info(f"✅ Imported {len(records)} records")
                self.stdout.write(self.style.SUCCESS(f"Imported {len(records)} new records"))
                return True
            else:
                logging.info("No new records to import")
                self.stdout.write(self.style.SUCCESS("No new records to import"))
                return True

        except Exception as e:
            logging.error(f"❌ Error importing to Django: {e}")
            self.stdout.write(self.style.ERROR(f"Import error: {str(e)}"))
            return False

    def cleanup_temp_files(self):
        """Clean up temporary files"""
        try:
            patterns = [
                os.path.join(settings.BASE_DIR, '*.xlsx'),
                os.path.join(settings.BASE_DIR, '*.csv'),
                os.path.join(settings.BASE_DIR, '*.db'),
                os.path.join(settings.BASE_DIR, '*.py'),
            ]
            
            for pattern in patterns:
                for file in glob.glob(pattern):
                    try:
                        if not any(f in file for f in ['RI.xlsx', 'RI.csv', 'RI.db', 'RI_News.docx']):
                            os.remove(file)
                            logging.info(f"Removed temporary file: {file}")
                    except Exception as e:
                        logging.warning(f"Could not remove {file}: {e}")
            
            self.stdout.write(self.style.SUCCESS("Cleaned up temporary files"))
            return True
        except Exception as e:
            logging.error(f"Error during cleanup: {e}")
            return False

    def determine_agency(self, source_file):
        """Determine agency from source filename"""
        if not source_file:
            return "Other"
        
        source_lower = str(source_file).lower()
        
        agency_mapping = {
            'ema': 'EMA',
            'ec': 'European Commission',
            'fda': 'FDA',
            'who': 'WHO',
            'mhra': 'MHRA',
            'ich': 'ICH',
            'gmp': 'GMP',
            'hma': 'HMA',
            'at': 'Austria',
            'be': 'Belgium',
            'cy': 'Cyprus',
            'de': 'Germany',
            'dk': 'Denmark',
            'fi': 'Finland',
            'ie': 'Ireland',
            'lux': 'Luxembourg',
            'mt': 'Malta',
            'no': 'Norway',
            'se': 'Sweden',
            'ch': 'Switzerland',
        }
        
        for key, agency in agency_mapping.items():
            if key in source_lower:
                return agency
        
        return "Other"

    def determine_category(self, source_file):
        """Determine category from source filename"""
        if not source_file:
            return "General"
        
        source_lower = str(source_file).lower()
        
        if 'policy' in source_lower:
            return "Policy"
        elif 'medical' in source_lower:
            return "Medical"
        elif 'update' in source_lower:
            return "Update"
        elif 'alert' in source_lower:
            return "Alert"
        elif 'legislation' in source_lower:
            return "Legislation"
        return "General"
    
    def parse_date(self, date_str):
        if pd.isna(date_str) or not str(date_str).strip():
            return None
        try:
            # Try strict parsing first
            return datetime.strptime(str(date_str), '%Y-%m-%d').date()
        except ValueError:
            try:
                # Fallback to flexible parsing
                return pd.to_datetime(date_str, errors='coerce').date()
            except:
                return None
